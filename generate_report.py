"""Generate DOCX: Existing Systems - Document Q&A Chatbots"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:val'):'clear', qn('w:color'):'auto', qn('w:fill'):color}))

def styled_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(255,255,255)
        shade(c, '1A1A2E')
    for ri,rd in enumerate(rows):
        for ci,ct in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(ct)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 0: shade(c, 'F0F4F8')
    return t

def heading(doc, txt, lvl=1):
    h = doc.add_heading(txt, level=lvl)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A,0x1A,0x2E)
    return h

def normal(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.size = Pt(10)
    return p

def bullet(doc, label, txt):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(label); r.bold = True; r.font.size = Pt(10)
    r = p.add_run(txt); r.font.size = Pt(10)

doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.54); s.bottom_margin=Cm(2.54)
    s.left_margin=Cm(2.54); s.right_margin=Cm(2.54)

t = doc.add_heading('Existing Systems - Document Q&A Chatbots', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in t.runs: r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E); r.font.size=Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('A Research Comparison of Available Solutions'); r.font.size=Pt(12)
r.font.color.rgb=RGBColor(0x66,0x66,0x66); r.italic=True
doc.add_paragraph()

# ===================== 1. ChatPDF =====================
heading(doc, '1. ChatPDF')
normal(doc, 'ChatPDF is a cloud-based web application that allows users to upload a PDF document and ask natural language questions about its content.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Cloud-based web app'],
    ['How it works','Upload PDF, ask questions, get answers'],
    ['Retrieval Method','Semantic embeddings + chunking'],
    ['LLM Used','GPT-based'],
    ['Supported Formats','PDF only'],
    ['Free Tier','Yes - 2 PDFs/day'],
    ['Paid Plan','$5/month'],
    ['Limitations','Single document at a time, no multi-format support, no export, no flowchart generation'],
])
doc.add_paragraph()

# ===================== 2. NotebookLM =====================
heading(doc, '2. Google NotebookLM')
normal(doc, 'Google NotebookLM is a cloud-based research tool that lets users create notebooks with up to 50 sources and chat with them using Google Gemini AI.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Cloud-based web app (Google)'],
    ['How it works','Create notebooks with up to 50 sources, chat with them'],
    ['Retrieval Method','Google semantic search'],
    ['LLM Used','Gemini'],
    ['Supported Formats','PDF, Google Docs, text, URLs, YouTube'],
    ['Free Tier','Yes - unlimited notebooks'],
    ['Paid Plan','NotebookLM Plus $19.99/month'],
    ['Unique Feature','Audio Overview - turns documents into podcast-style summaries'],
    ['Limitations','No Word/PowerPoint/Excel support, no API, cloud-only'],
])
doc.add_paragraph()

# ===================== 3. AnythingLLM =====================
heading(doc, '3. AnythingLLM')
normal(doc, 'AnythingLLM is an open-source tool that bundles document ingestion, embeddings, vector storage, and chat into a single application. It supports over 30 LLM providers.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Open-source, self-hosted (Docker)'],
    ['How it works','Docker deployment, ingest documents, chat locally'],
    ['Retrieval Method','Vector embeddings (configurable)'],
    ['LLM Used','30+ providers (Ollama, OpenAI, Groq, etc.)'],
    ['Supported Formats','PDF, DOCX, TXT, URLs, YouTube, GitHub, Confluence'],
    ['Free Tier','Yes - fully open-source'],
    ['Paid Plan','Cloud tier available'],
    ['Unique Feature','Multi-user support, workspace isolation, fully offline capable'],
    ['Limitations','Complex setup, retrieval quality depends on configuration'],
])
doc.add_paragraph()

# ===================== 4. Docora =====================
heading(doc, '4. Docora')
normal(doc, 'Docora is a desktop application that runs locally. It indexes PDFs, Word documents, PowerPoint presentations, and Excel spreadsheets from a selected folder.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Desktop app (Mac-first, Windows coming)'],
    ['How it works','Point at a folder, index all docs, chat'],
    ['Retrieval Method','Hybrid: Semantic embeddings + BM25 keyword matching + reranking'],
    ['LLM Used','OpenAI / Google / VoyageAI (user API key)'],
    ['Supported Formats','PDF, DOCX, PPTX, XLSX'],
    ['Free Tier','Yes (limited)'],
    ['Unique Feature','Source citations with page numbers, files stay local'],
    ['Limitations','Mac-first, requires API key'],
])
doc.add_paragraph()

# ===================== 5. PrivateGPT =====================
heading(doc, '5. PrivateGPT')
normal(doc, 'PrivateGPT is an open-source tool designed for fully offline document AI. It runs entirely on the users machine with no internet connection required.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Open-source, fully offline (Python)'],
    ['How it works','Python-based, runs entirely on your machine'],
    ['Retrieval Method','Local embeddings + local LLM'],
    ['LLM Used','Ollama / llama.cpp (local models)'],
    ['Supported Formats','PDF, DOCX, TXT'],
    ['Free Tier','Yes - fully open-source'],
    ['Unique Feature','100% air-gap capable, no internet needed'],
    ['Limitations','Requires 16GB+ RAM, complex setup, lower AI quality than cloud models'],
])
doc.add_paragraph()

# ===================== 6. Khoj =====================
heading(doc, '6. Khoj')
normal(doc, 'Khoj is an open-source personal AI that can be self-hosted or used via cloud. It allows users to chat with their documents and the web.')
styled_table(doc, ['Feature','Details'], [
    ['Type','Open-source personal AI (Docker / cloud)'],
    ['How it works','Self-hosted or cloud, chat with documents + web'],
    ['Retrieval Method','Semantic search'],
    ['LLM Used','Multiple providers'],
    ['Supported Formats','PDF, markdown, org-mode, plain text, GitHub, Google Drive'],
    ['Free Tier','Yes (self-hosted) / $8/month cloud'],
    ['Unique Feature','Calendar integration, personal knowledge base'],
    ['Limitations','Limited file format support'],
])
doc.add_paragraph()

# ===================== COMPARISON TABLE =====================
heading(doc, 'Comparative Analysis', 1)
normal(doc, 'The following table provides a side-by-side comparison of all existing systems with key features relevant to document Q&A chatbots.')
doc.add_paragraph()
styled_table(doc,
    ['Feature','ChatPDF','NotebookLM','AnythingLLM','Docora','PrivateGPT','Khoj','DocChat (Ours)'],
    [
        ['Cost','$5/mo','Free','Free (OSS)','Paid','Free (OSS)','Free / $8mo','Free (OSS)'],
        ['PDF Support','Yes','Yes','Yes','Yes','Yes','Yes','Yes'],
        ['DOCX Support','No','No','Yes','Yes','Yes','No','Yes'],
        ['TXT Support','No','Yes','Yes','No','Yes','Yes','Yes'],
        ['Offline Mode','No','No','Yes','No','Yes','No','No'],
        ['Multi-document','No','Yes (50)','Yes','Yes','No','No','No'],
        ['Flowchart Generation','No','No','No','No','No','No','Yes'],
        ['Mobile Responsive','No','Yes','No','No','No','No','Yes'],
        ['Custom Model Choice','No','No','Yes','Yes','Yes','Yes','Yes'],
        ['Setup Difficulty','None','None','Medium','Easy','Hard','Medium','Easy'],
        ['Self-hosted','No','No','Yes','No','Yes','Yes','Yes'],
        ['Source Citations','No','Yes','No','Yes','No','No','No'],
    ]
)
doc.add_paragraph()

# ===================== UNIQUE ADVANTAGES =====================
heading(doc, 'Unique Advantages of Our System (DocChat)', 1)
normal(doc, 'Based on the comparative analysis, our proposed system DocChat offers the following unique advantages over existing solutions:')
doc.add_paragraph()

bullet(doc, 'Automatic Flowchart Generation: ', 'No other tool in the market generates interactive Mermaid flowcharts from document content. Our system extracts process steps and renders them as SVG diagrams automatically.')
bullet(doc, 'Zero Cost: ', 'The entire stack is open-source and free. No paid API subscriptions, no paid hosting. Groq free tier provides sufficient API access for development and testing.')
bullet(doc, 'Multi-Format Support: ', 'Supports PDF, DOCX, and TXT in a single interface. Most competitors support only PDF.')
bullet(doc, 'Mobile-Responsive Design: ', 'Fully responsive with a slide-in drawer sidebar for mobile. Most document chat tools are desktop-only.')
bullet(doc, 'Adjustable Detail Level: ', 'Users can control retrieval depth via a slider, balancing between concise and detailed answers.')
bullet(doc, 'No-Hallucination Policy: ', 'The system prompt strictly enforces answering only from retrieved context, with explicit statements when information is not available in the document.')
bullet(doc, 'Multi-Model Support: ', 'Users can switch between GPT OSS 120B, GPT OSS 20B, and Qwen 3.6 27B from a dropdown.')
bullet(doc, 'Easy Setup: ', 'Single command installation via pip. No Docker, no complex configuration required.')
doc.add_paragraph()

# ===================== REFERENCES =====================
heading(doc, 'References', 1)
refs = [
    'ChatPDF - https://www.chatpdf.com',
    'Google NotebookLM - https://notebooklm.google.com',
    'AnythingLLM - https://github.com/Mintplex-Labs/anything-llm',
    'Docora - https://docora.dev',
    'PrivateGPT - https://github.com/zylon-ai/private-gpt',
    'Khoj - https://khoj.dev',
    'Groq API Documentation - https://console.groq.com/docs/models',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f'[{i}] {ref}'); r.font.size = Pt(9)

# ===================== SAVE =====================
output = 'Existing_Systems_Document_QA_Chatbots.docx'
doc.save(output)
print(f'Report saved: {output}')

